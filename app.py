from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from openai import OpenAI

app = Flask(__name__)
CORS(app)

# Новый синтаксис OpenAI >= 1.0.0
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

def generate_with_gpt(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a CLIL methodology expert. Generate a list of classroom tasks for school teachers using the CLIL (Content and Language Integrated Learning) approach. Each task should:\n\n- follow one of the 15 formats (matching, short_answer, true_false, multiple_choice, fill_in_the_blank, ordering, open_ended, dialogue_completion, categorization, matching_definitions, gap_fill, table_completion, problem_solving, scenario_based, case_analysis);\n- integrate subject content, target language functions (e.g. define, explain, describe), and national or cultural values;\n- be written clearly and fully in one paragraph per task;\n- **do not number inner elements** inside the task body.\n\nUse the format:\nTask 1: ...\nTask 2: ..."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )
    return response.choices[0].message.content

@app.route("/generate_gpt", methods=["POST"])
def generate_gpt():
    data = request.json
    topic = data["topic"]

    # CLIL принциптері бойынша, пән – Информатика
    prompt = (
        f"Generate one CLIL-based task set for the topic '{topic}' in the subject 'Informatics'.\n"
        f"Use three integrated parts:\n"
        f"📖 Reading – A task that asks the learner to read and understand the topic.\n"
        f"✍️ Writing – A task that asks the learner to write or express in written form.\n"
        f"🗣️ Speaking – A task that asks the learner to explain or describe the topic verbally.\n\n"
        f"Each part should be 1–2 sentences long and clearly related to informatics.\n"
        f"Return in JSON format like:\n"
        f'{{"reading": "...", "writing": "...", "speaking": "..."}}'
    )

    result_text = generate_with_gpt(prompt)

    # JSON форматына парсинг
    import json
    try:
        result_json = json.loads(result_text)
    except json.JSONDecodeError:
        result_json = {
            "reading": f"Осы тақырыпқа қатысты мәтінді оқып, негізгі идеяларын анықтаңыз.",
            "writing": f"'{topic}' тақырыбы бойынша өз ойыңызды жазбаша түрде білдіріңіз.",
            "speaking": f"'{topic}' тақырыбын сыныптастарыңызбен талқылап, түсіндіріңіз."
        }

    return jsonify(result_json)


@app.route("/download_docx", methods=["POST"])
def download_docx():
    data = request.json
    tasks = data.get("tasks", [])
    doc = Document()
    doc.add_heading('CLIL Tasks', level=1)
    for i, task in enumerate(tasks, start=1):
        doc.add_paragraph(f"{i}. {task}")
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return send_file(temp_file.name, as_attachment=True, download_name="clil_tasks.docx")

@app.route("/download_pdf", methods=["POST"])
def download_pdf():
    data = request.json
    tasks = data.get("tasks", [])

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    c = canvas.Canvas(temp_file.name, pagesize=letter)

    width, height = letter
    y = height - 60

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "CLIL Tasks")
    y -= 30

    c.setFont("Helvetica", 12)

    for i, task in enumerate(tasks, start=1):
        wrapped_text = f"{i}. {task}"
        text = c.beginText(50, y)
        text.setFont("Helvetica", 12)

        # Автоматически разбиваем на строки по 100 символов
        max_chars = 100
        lines = [wrapped_text[j:j+max_chars] for j in range(0, len(wrapped_text), max_chars)]
        for line in lines:
            text.textLine(line)
            y -= 15
            if y < 60:
                c.drawText(text)
                c.showPage()
                y = height - 60
                text = c.beginText(50, y)

        c.drawText(text)
        y -= 15

    c.save()
    return send_file(temp_file.name, as_attachment=True, download_name="clil_tasks.pdf")
